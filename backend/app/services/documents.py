"""文档服务：模板变量回填、Word 导出、SBOM 生成、整改报告。"""
import io
import re
from datetime import datetime, timezone

from docx import Document as Docx
from sqlalchemy.orm import Session

from ..models import Component, ControlItem, AssessmentAnswer, OrgNode


def render_template(body_html: str, fields: dict) -> str:
    """把 {{key}} 占位符替换为 fields 值。"""
    def repl(m):
        key = m.group(1).strip()
        return str(fields.get(key, m.group(0)))
    return re.sub(r"\{\{\s*([\w.]+)\s*\}\}", repl, body_html or "")


def html_to_docx_bytes(title: str, content_html: str) -> bytes:
    """极简 HTML→Word：剥离标签按段落写入（满足交付导出，不追求富样式）。"""
    doc = Docx()
    doc.add_heading(title or "文档", level=0)
    # 按块级标签切段
    blocks = re.split(r"</(?:p|div|h[1-6]|li|tr)>", content_html or "")
    for block in blocks:
        # 标题识别
        heading = re.search(r"<h([1-6])[^>]*>", block)
        text = re.sub(r"<[^>]+>", "", block)
        text = re.sub(r"&nbsp;", " ", text).strip()
        if not text:
            continue
        if heading:
            doc.add_heading(text, level=min(int(heading.group(1)), 4))
        else:
            doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_sbom(db: Session, product_node_id: int) -> dict:
    """从组件表生成 CycloneDX 1.5 JSON。"""
    comps = db.query(Component).filter(Component.product_node_id == product_node_id).all()
    node = db.get(OrgNode, product_node_id)
    return {
        "bomFormat": "CycloneDX",
        "specVersion": "1.5",
        "version": 1,
        "metadata": {
            "timestamp": datetime.now(timezone.utc).isoformat(),
            "component": {"type": "application", "name": node.name if node else "product"},
        },
        "components": [
            {
                "type": "library",
                "name": c.name,
                "version": c.version,
                "purl": c.purl or f"pkg:generic/{c.name}@{c.version}",
                "licenses": [{"license": {"id": c.license}}] if c.license else [],
            }
            for c in comps
        ],
    }


def build_remediation_report(db: Session, assessment_id: int) -> list[dict]:
    """对比每条答案等级与 ControlItem.remediation 规则，输出差距+建议+推荐工具。"""
    answers = db.query(AssessmentAnswer).filter(
        AssessmentAnswer.assessment_id == assessment_id).all()
    out = []
    for a in answers:
        item = db.get(ControlItem, a.item_id)
        if not item:
            continue
        for rule in (item.remediation or []):
            if a.level <= rule.get("level_threshold", 2):
                out.append({
                    "item_code": item.code,
                    "item_title": item.title,
                    "cra_ref": item.cra_ref,
                    "current_level": a.level,
                    "gap": rule.get("gap_desc", ""),
                    "advice": rule.get("advice", ""),
                    "recommended_tools": rule.get("recommended_tools", []),
                })
                break
    return out
